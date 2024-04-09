import pandas as pd
import openpyxl
import Funciones
from Funciones import funciones_Word
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import messagebox

# Función para ejecutar el código proporcionado
def ejecutar_codigo(archivo_path, nombre_hoja):
    #abrimos el archivo selecionado
    wb = openpyxl.load_workbook(archivo_path)

    #ubicarnos en la hoja actual
    ws = wb[nombre_hoja]

    #Quitar el filtro a las tablas con filtros
    ws.auto_filter.ref = None

    #Guardar los cambios
    wb.save(archivo_path)

    #Leer cada dato del archivo excel proporcionado por el cliente y del formato
    data_cli = pd.ExcelFile(archivo_path)

    df_relationship = None  # Asegúrate de definir df_relationship aquí o en caso de error
    for sheet_name_cli in data_cli.sheet_names:
        if sheet_name_cli == nombre_hoja:
            #Leer los dato de la hoja actual
            df_hoja_actual = pd.read_excel(archivo_path, sheet_name=sheet_name_cli, header=None)

            # Desactivar los filtros si están activos
            df_hoja_actual = df_hoja_actual.iloc[1:]# Omitir la fila con los filtros activados

            # Encontrar la fila que contiene la mayoría de valores no nulos
            fila_encabezados = df_hoja_actual.notna().sum(axis=1).idxmax()

            # Volver a leer los datos de la hoja actual, especificando la fila de encabezados
            df_hoja_actual = pd.read_excel(archivo_path, sheet_name=sheet_name_cli, header=fila_encabezados)

            # Eliminar filas completamente vacías
            df_hoja_actual = df_hoja_actual.dropna(how='all', axis=1).dropna(how='all')

            # Mostrar una muestra de los datos de la hoja actual
            print(f"Datos de la hoja '{sheet_name_cli}':")
            print(df_hoja_actual.head())  # Muestra los primeros 5 registros por defecto
            print("\n")  # Agrega una línea en blanco entre cada hoja

            #Almacenar los datos de la hoja actual en el diaccionario
            df_relationship = df_hoja_actual
            break
    #generar_informe_errores(df_relationship)
    return df_relationship

def generar_informe_errores(df_relationship, datos_seleccionados_lista):
    # Crear un nuevo documento de Word para el informe
    doc_informe = Document()

    def agregar_titulo_informe(titulo):
        p = doc_informe.add_paragraph()
        p.add_run(titulo).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_informe.add_paragraph()

    def agregar_tabla_informe(data):
        table = doc_informe.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Descripción del Error'
        hdr_cells[1].text = 'Detalle'

        for descripcion, detalle in data:
            row_cells = table.add_row().cells
            row_cells[0].text = descripcion
            row_cells[1].text = detalle

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Inicializar una lista para almacenar los errores encontrados
    errores = []

    # Recorrer los datos seleccionados almacenados en datos_seleccionados_lista
    for seleccion in datos_seleccionados_lista:
        # Dividir la selección en partes utilizando algún método específico
        # Supongamos que la selección está en formato 'CÓDIGO RELATIONSHIP|DESCRIPCIÓN'
        codigo_relacion, descripcion = seleccion.split('|')

        # Verificar si el código de relación no es numérico
        if not codigo_relacion.isdigit():
            errores.append(("Código de relación no numérico", f"Valor encontrado: {codigo_relacion}"))

        # Verificar si la descripción contiene la palabra 'titular' y el código de relación es diferente de 1
        if 'titular' in descripcion.lower() and codigo_relacion != '1':
            errores.append(("Código de relación incorrecto para descripción 'titular', titular siempre va con 1", f"Código de relación: {codigo_relacion}, Descripción: {descripcion}"))

    # Si se encontraron errores, generar el informe
    if errores:
        agregar_titulo_informe("Informe de Errores")

        # Agregar la información resaltada al informe
        doc_informe.add_paragraph("Información resaltada:")
        #doc_informe.add_paragraph(texto_resaltado)

        # Agregar la tabla de errores al informe
        agregar_tabla_informe(errores)

        # Guardar el informe
        doc_informe.save('Informe_errores.docx')
        messagebox.showinfo("Éxito", "Documento Word generado con éxito")
    else:
        messagebox.showinfo("Éxito", "No se encontraron errores.")



    
