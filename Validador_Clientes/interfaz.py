import pandas as pd
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Función para ejecutar el código proporcionado
def ejecutar_codigo(archivo_path, nombre_hoja):
    # Crear un nuevo documento de Word para el informe
    doc_informe = Document()

    # Lista para almacenar los errores encontrados
    errores = []

    wb = openpyxl.load_workbook(archivo_path)
    ws = wb[nombre_hoja]
    ws.auto_filter.ref = None
    wb.save(archivo_path)

    data_cli = pd.ExcelFile(archivo_path)

    for sheet_name_cli in data_cli.sheet_names:
        if sheet_name_cli == nombre_hoja:
            df_hoja_actual = pd.read_excel(archivo_path, sheet_name=sheet_name_cli, header=None)
            df_hoja_actual = df_hoja_actual.iloc[1:]
            fila_encabezados = df_hoja_actual.notna().sum(axis=1).idxmax()
            df_hoja_actual = pd.read_excel(archivo_path, sheet_name=sheet_name_cli, header=fila_encabezados)
            df_hoja_actual = df_hoja_actual.dropna(how='all', axis=1).dropna(how='all')
            df_relationship = df_hoja_actual
            break

    def agregar_titulo_informe(titulo):
        p = doc_informe.add_paragraph()
        p.add_run(titulo).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_informe.add_paragraph()

    def agregar_tabla_informe(data):
        table = doc_informe.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Descripción del Error'
        hdr_cells[1].text = 'Detalle'

        for descripcion, detalle in data:
            row_cells = table.add_row().cells
            row_cells[0].text = descripcion
            row_cells[1].text = detalle

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    errores = []
    for _, row in df_relationship.iterrows():
        codigo_relacion = row['CÓDIGO RELATIONSHIP']
        descripcion = row['DESCRIPCIÓN']

        if not str(codigo_relacion).isdigit():
            errores.append(("Código de relación no numérico", f"Valor encontrado: {codigo_relacion}"))

        if 'titular' in descripcion.lower():
            if codigo_relacion != 1:
                errores.append(("Código de relación incorrecto para descripción 'titular', titular siempre va con 1", f"Código de relación: {codigo_relacion}, Descripción: {descripcion}"))

    if errores:
        agregar_titulo_informe("Informe de Errores")
        agregar_tabla_informe(errores)
        doc_informe.save('Informe_errores.docx')
        print("Documento Word generado con éxito")
    else:
        print("No se encontraron errores.")
