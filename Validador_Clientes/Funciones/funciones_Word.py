from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Función para agregar un título al informe
def agregar_titulo_informe(titulo, doc_informe):
    p = doc_informe.add_paragraph()
    p.add_run(titulo).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_informe.add_paragraph() # Agregar un espacio en blanco después del título


# Función para agregar una tabla al informe
def agregar_tabla_informe(data, doc_informe):
    table = doc_informe.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Encabezados de la tabla
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Descripción del Error'
    hdr_cells[1].text = 'Detalle'

    # Agregar datos a la tabla
    for descripcion, detalle in data:
        row_cells = table.add_row().cells
        row_cells[0].text = descripcion
        row_cells[1].text = detalle

    # Alinear verticalmente el contenido de la tabla    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)